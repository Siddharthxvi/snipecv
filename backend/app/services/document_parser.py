import io
import fitz # PyMuPDF
import docx
from fastapi import HTTPException

class DocumentParser:
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX document: {str(e)}")

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("latin-1")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to decode text file: {str(e)}")

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> str:
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return cls.parse_pdf(file_bytes)
        elif ext == "docx":
            return cls.parse_docx(file_bytes)
        elif ext == "txt":
            return cls.parse_txt(file_bytes)
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file extension .{ext}. Supported formats are: .pdf, .docx, .txt"
            )
